from collections.abc import Generator
from typing import Any

from dify_plugin import Tool
from dify_plugin.entities.tool import ToolInvokeMessage

import os
import re
from typing import List, Optional, Tuple
from urllib.parse import urljoin, urlparse
import requests
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.ns import qn
from io import BytesIO
from PIL import Image
import traceback


class MarkdownParser:
    """解析 Markdown 内容"""

    def __init__(self, content: str):
        self.content = content
        self.lines = content.split('\n')

    def _is_table_row(self, line: str) -> bool:
        """判断是否是表格行"""
        return '|' in line and line.strip().startswith('|') or line.count('|') >= 2

    def _is_table_separator(self, line: str) -> bool:
        """判断是否是表格分隔行"""
        stripped = line.strip()
        if not self._is_table_row(stripped):
            return False
        # 分隔行应该只包含 |, -, : 和空格
        cleaned = stripped.replace('|', '').replace('-', '').replace(':', '').replace(' ', '')
        return len(cleaned) == 0

    def _parse_table_row(self, line: str) -> List[str]:
        """解析表格行"""
        # 移除首尾的 |
        line = line.strip()
        if line.startswith('|'):
            line = line[1:]
        if line.endswith('|'):
            line = line[:-1]
        # 分割单元格
        cells = [cell.strip() for cell in line.split('|')]
        return cells

    def _parse_table(self, start_idx: int) -> Tuple[dict, int]:
        """解析表格，返回表格数据和下一行索引"""
        table_lines = []
        i = start_idx

        # 收集所有表格行
        while i < len(self.lines) and self._is_table_row(self.lines[i]):
            table_lines.append(self.lines[i])
            i += 1

        if len(table_lines) < 2:
            # 至少需要表头和分隔行
            return None, start_idx + 1

        # 解析表头
        headers = self._parse_table_row(table_lines[0])

        # 检查第二行是否是分隔行
        if not self._is_table_separator(table_lines[1]):
            return None, start_idx + 1

        # 解析数据行
        rows = []
        for line in table_lines[2:]:
            if line.strip():
                rows.append(self._parse_table_row(line))

        table_data = {
            'type': 'table',
            'headers': headers,
            'rows': rows
        }

        return table_data, i

    def _is_list_item(self, line: str) -> Tuple[bool, Optional[str], Optional[str]]:
        """
        判断是否是列表项，返回 (是否列表, 列表类型, 列表内容)
        列表类型: 'unordered' 或 'ordered'
        """
        stripped = line.strip()

        # 无序列表: -, *, +
        unordered_match = re.match(r'^([-*+])\s+(.+)$', stripped)
        if unordered_match:
            return True, 'unordered', unordered_match.group(2)

        # 有序列表: 1. 2. 等
        ordered_match = re.match(r'^\d+\.\s+(.+)$', stripped)
        if ordered_match:
            return True, 'ordered', ordered_match.group(1)

        return False, None, None

    def _parse_list(self, start_idx: int) -> Tuple[dict, int]:
        """解析列表，返回列表数据和下一行索引"""
        list_items = []
        list_type = None
        i = start_idx

        while i < len(self.lines):
            line = self.lines[i]
            is_list, item_type, content = self._is_list_item(line)

            if not is_list:
                # 检查是否是空行（列表可能继续）
                if not line.strip():
                    # 如果下一行也是空行或不是列表，则结束列表
                    if i + 1 >= len(self.lines) or not self.lines[i + 1].strip():
                        break
                    # 检查下一行是否是列表
                    next_is_list, _, _ = self._is_list_item(self.lines[i + 1])
                    if not next_is_list:
                        break
                    i += 1
                    continue
                else:
                    break

            # 确定列表类型（第一次遇到时）
            if list_type is None:
                list_type = item_type

            # 只处理相同类型的列表项
            if item_type == list_type:
                list_items.append({
                    'content': content,
                    'type': item_type
                })
            else:
                break

            i += 1

        if not list_items:
            return None, start_idx + 1

        list_data = {
            'type': 'list',
            'list_type': list_type,
            'items': list_items
        }

        return list_data, i

    def _parse_inline_formatting(self, text: str) -> List[dict]:
        """
        解析行内格式（加粗、斜体等）
        返回格式化的文本片段列表
        每个片段格式: {'text': str, 'bold': bool, 'italic': bool}
        """
        parts = []
        # 匹配加粗和斜体的模式
        # 优先级：加粗 > 斜体
        pattern = r'(\*\*\*([^*]+)\*\*\*|\*\*([^*]+)\*\*|\*([^*]+)\*|__([^_]+)__|_([^_]+)_)'

        last_pos = 0
        for match in re.finditer(pattern, text):
            # 添加匹配前的普通文本
            if match.start() > last_pos:
                parts.append({
                    'text': text[last_pos:match.start()],
                    'bold': False,
                    'italic': False
                })

            # 处理匹配的格式
            full_match = match.group(0)
            # 加粗斜体: ***text***
            if full_match.startswith('***') and full_match.endswith('***'):
                content = match.group(2)
                parts.append({
                    'text': content,
                    'bold': True,
                    'italic': True
                })
            # 加粗: **text** 或 __text__
            elif full_match.startswith('**') and full_match.endswith('**'):
                content = match.group(3)
                parts.append({
                    'text': content,
                    'bold': True,
                    'italic': False
                })
            elif full_match.startswith('__') and full_match.endswith('__'):
                content = match.group(5)
                parts.append({
                    'text': content,
                    'bold': True,
                    'italic': False
                })
            # 斜体: *text* 或 _text_
            elif full_match.startswith('*') and full_match.endswith('*'):
                content = match.group(4)
                parts.append({
                    'text': content,
                    'bold': False,
                    'italic': True
                })
            elif full_match.startswith('_') and full_match.endswith('_'):
                content = match.group(6)
                parts.append({
                    'text': content,
                    'bold': False,
                    'italic': True
                })

            last_pos = match.end()

        # 添加剩余的普通文本
        if last_pos < len(text):
            parts.append({
                'text': text[last_pos:],
                'bold': False,
                'italic': False
            })

        # 如果没有匹配到任何格式，返回原始文本
        if not parts:
            parts.append({
                'text': text,
                'bold': False,
                'italic': False
            })

        return parts

    def parse(self) -> List[dict]:
        """解析 Markdown 内容为结构化数据"""
        elements = []
        i = 0
        in_code_block = False
        code_lines = []
        code_lang = ''

        while i < len(self.lines):
            line = self.lines[i]

            # 处理代码块
            if line.strip().startswith('```'):
                if not in_code_block:
                    in_code_block = True
                    code_lang = line.strip()[3:].strip()
                    code_lines = []
                else:
                    in_code_block = False
                    elements.append({
                        'type': 'code',
                        'content': '\n'.join(code_lines),
                        'language': code_lang
                    })
                    code_lines = []
                    code_lang = ''
                i += 1
                continue

            if in_code_block:
                code_lines.append(line)
                i += 1
                continue

            # 处理标题
            heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if heading_match:
                level = len(heading_match.group(1))
                text = heading_match.group(2).strip()
                elements.append({
                    'type': 'heading',
                    'level': level,
                    'content': text
                })
                i += 1
                continue

            # 处理表格
            if self._is_table_row(line):
                table_data, next_i = self._parse_table(i)
                if table_data:
                    elements.append(table_data)
                    i = next_i
                    continue

            # 处理列表
            is_list, _, _ = self._is_list_item(line)
            if is_list:
                list_data, next_i = self._parse_list(i)
                if list_data:
                    elements.append(list_data)
                    i = next_i
                    continue

            # 处理独立的图片
            img_match = re.match(r'^!\[([^\]]*)\]\(([^\)]+)\)\s*$', line)
            if img_match:
                alt_text = img_match.group(1)
                img_url = img_match.group(2)
                elements.append({
                    'type': 'image',
                    'alt': alt_text,
                    'url': img_url
                })
                i += 1
                continue

            # 处理空行
            if not line.strip():
                i += 1
                continue

            # 处理普通段落（可能包含行内图片和格式）
            paragraph_lines = []
            while i < len(self.lines) and self.lines[i].strip():
                if self.lines[i].startswith('#') or self.lines[i].strip().startswith('```'):
                    break
                # 检查是否是表格行
                if self._is_table_row(self.lines[i]):
                    break
                # 检查是否是列表项
                is_list_item, _, _ = self._is_list_item(self.lines[i])
                if is_list_item:
                    break
                # 检查是否是独立图片行
                if re.match(r'^!\[([^\]]*)\]\(([^\)]+)\)\s*$', self.lines[i]):
                    break
                paragraph_lines.append(self.lines[i])
                i += 1

            if paragraph_lines:
                paragraph_text = ' '.join(paragraph_lines)
                # 检查段落中是否包含图片
                img_pattern = r'!\[([^\]]*)\]\(([^\)]+)\)'
                if re.search(img_pattern, paragraph_text):
                    # 拆分文本和图片
                    parts = re.split(img_pattern, paragraph_text)
                    for idx, part in enumerate(parts):
                        if idx % 3 == 0 and part.strip():  # 文本部分
                            # 解析行内格式
                            formatted_parts = self._parse_inline_formatting(part.strip())
                            elements.append({
                                'type': 'paragraph',
                                'content': part.strip(),
                                'formatted_parts': formatted_parts
                            })
                        elif idx % 3 == 2:  # 图片 URL
                            alt_text = parts[idx - 1]
                            elements.append({
                                'type': 'image',
                                'alt': alt_text,
                                'url': part
                            })
                else:
                    # 解析行内格式
                    formatted_parts = self._parse_inline_formatting(paragraph_text)
                    elements.append({
                        'type': 'paragraph',
                        'content': paragraph_text,
                        'formatted_parts': formatted_parts
                    })

        return elements


class DocxConverter:
    """将解析的 Markdown 转换为 DOCX"""

    # 字号定义（磅值）
    FONT_SIZE_3 = Pt(16)  # 三号字
    FONT_SIZE_4 = Pt(14)  # 四号字

    def __init__(self, base_url: Optional[str] = None, image_width: float = 5.0):
        """
        初始化转换器

        Args:
            base_url: 基础 URL，用于解析相对路径的图片
            image_width: 图片宽度（英寸）
        """
        self.document = Document()
        self.base_url = base_url
        self.image_width = image_width

    def set_song_font(self, run):
        """设置宋体字体"""
        run.font.name = '宋体'
        # 设置中文字体
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    def add_heading(self, text: str, level: int):
        """添加标题"""
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(text)

        # 一级标题使用三号字，其他使用四号字
        if level == 1:
            run.font.size = self.FONT_SIZE_3
        else:
            run.font.size = self.FONT_SIZE_4

        run.font.bold = True
        self.set_song_font(run)

        # 设置行间距
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    def add_paragraph(self, text: str, formatted_parts: Optional[List[dict]] = None):
        """添加段落"""
        paragraph = self.document.add_paragraph()

        # 如果有格式化的部分，使用格式化文本
        if formatted_parts:
            for part in formatted_parts:
                run = paragraph.add_run(part['text'])
                run.font.size = self.FONT_SIZE_4
                self.set_song_font(run)
                if part.get('bold'):
                    run.font.bold = True
                if part.get('italic'):
                    run.font.italic = True
        else:
            # 普通文本
            run = paragraph.add_run(text)
            run.font.size = self.FONT_SIZE_4
            self.set_song_font(run)

        # 设置段落格式
        paragraph.paragraph_format.first_line_indent = Pt(28)  # 2字符缩进（约28磅）
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    def add_code(self, code_text: str, language: str = ''):
        """添加代码块"""
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(code_text)

        # 代码使用等宽字体，较小字号
        run.font.name = 'Consolas'
        run.font.size = Pt(10)

        # 设置背景色（浅灰色）
        paragraph.paragraph_format.left_indent = Pt(14)
        paragraph.paragraph_format.right_indent = Pt(14)

    def download_image(self, url: str) -> Optional[BytesIO]:
        """下载图片"""
        try:
            # 处理相对路径
            if self.base_url and not urlparse(url).scheme:
                url = urljoin(self.base_url, url)

            # 如果是本地文件路径
            if not urlparse(url).scheme or urlparse(url).scheme == 'file':
                # 去除 file:// 前缀
                local_path = url.replace('file://', '')
                if os.path.exists(local_path):
                    with open(local_path, 'rb') as f:
                        return BytesIO(f.read())
                else:
                    print(f"警告: 本地图片不存在: {local_path}")
                    return None

            # 下载网络图片
            response = requests.get(url, timeout=10)
            response.raise_for_status()
            return BytesIO(response.content)

        except Exception as e:
            print(f"警告: 无法下载图片 {url}: {str(e)}")
            return None

    def add_image(self, url: str, alt_text: str = ''):
        """添加图片"""
        image_data = self.download_image(url)

        if image_data:
            try:
                # 验证图片
                img = Image.open(image_data)
                img.verify()
                image_data.seek(0)  # 重置指针

                # 添加图片到文档
                paragraph = self.document.add_paragraph()
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = paragraph.add_run()
                run.add_picture(image_data, width=Inches(self.image_width))

                # 如果有描述文字，添加图注
                if alt_text:
                    caption = self.document.add_paragraph()
                    caption_run = caption.add_run(f"图: {alt_text}")
                    caption_run.font.size = Pt(10)
                    self.set_song_font(caption_run)
                    caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            except Exception as e:
                print(traceback.format_exc())
                print(f"警告: 无法插入图片 {url}: {str(e)}")
                # 添加替代文本
                self.add_paragraph(f"[图片: {alt_text or url}]")
        else:
            # 图片下载失败，添加替代文本
            self.add_paragraph(f"[图片: {alt_text or url}]")

    def add_table(self, headers: List[str], rows: List[List[str]]):
        """添加表格"""
        # 创建表格（行数 = 表头1行 + 数据行）
        table = self.document.add_table(rows=1 + len(rows), cols=len(headers))

        # 设置表格样式
        table.style = 'Light Grid Accent 1'

        # 填充表头
        header_cells = table.rows[0].cells
        for idx, header_text in enumerate(headers):
            cell = header_cells[idx]
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(header_text)

            # 设置表头格式：宋体四号字，加粗
            run.font.size = self.FONT_SIZE_4
            run.font.bold = True
            self.set_song_font(run)

            # 表头居中
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 填充数据行
        for row_idx, row_data in enumerate(rows):
            row_cells = table.rows[row_idx + 1].cells
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < len(row_cells):  # 防止列数不匹配
                    cell = row_cells[col_idx]
                    paragraph = cell.paragraphs[0]
                    run = paragraph.add_run(cell_text)

                    # 设置单元格格式：宋体四号字
                    run.font.size = self.FONT_SIZE_4
                    self.set_song_font(run)

        # 在表格后添加一个空段落，避免格式问题
        self.document.add_paragraph()

    def add_list(self, list_type: str, items: List[dict]):
        """添加列表（有序或无序）"""
        # 确定列表样式
        list_style = 'List Bullet' if list_type == 'unordered' else 'List Number'

        for item in items:
            # 创建列表段落
            paragraph = self.document.add_paragraph(style=list_style)

            # 解析列表项中的行内格式
            content = item.get('content', '')
            formatted_parts = self._parse_inline_formatting(content)

            # 清除默认创建的 run（如果有）
            if paragraph.runs:
                for run in paragraph.runs:
                    run.text = ''

            # 添加格式化的文本
            for part in formatted_parts:
                if part['text']:  # 只添加非空文本
                    run = paragraph.add_run(part['text'])
                    run.font.size = self.FONT_SIZE_4
                    self.set_song_font(run)
                    if part.get('bold'):
                        run.font.bold = True
                    if part.get('italic'):
                        run.font.italic = True

            # 设置行间距
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        # 在列表后添加一个空段落，避免格式问题
        self.document.add_paragraph()

    def _parse_inline_formatting(self, text: str) -> List[dict]:
        """
        解析行内格式（加粗、斜体等）
        返回格式化的文本片段列表
        """
        parts = []
        # 匹配加粗和斜体的模式
        # 优先级：加粗斜体 > 加粗 > 斜体
        pattern = r'(\*\*\*([^*]+)\*\*\*|\*\*([^*]+)\*\*|\*([^*]+)\*|__([^_]+)__|_([^_]+)_)'

        last_pos = 0
        for match in re.finditer(pattern, text):
            # 添加匹配前的普通文本
            if match.start() > last_pos:
                parts.append({
                    'text': text[last_pos:match.start()],
                    'bold': False,
                    'italic': False
                })

            # 处理匹配的格式
            full_match = match.group(0)
            # 加粗斜体: ***text***
            if full_match.startswith('***') and full_match.endswith('***'):
                content = match.group(2)
                parts.append({
                    'text': content,
                    'bold': True,
                    'italic': True
                })
            # 加粗: **text** 或 __text__
            elif full_match.startswith('**') and full_match.endswith('**'):
                content = match.group(3)
                parts.append({
                    'text': content,
                    'bold': True,
                    'italic': False
                })
            elif full_match.startswith('__') and full_match.endswith('__'):
                content = match.group(5)
                parts.append({
                    'text': content,
                    'bold': True,
                    'italic': False
                })
            # 斜体: *text* 或 _text_
            elif full_match.startswith('*') and full_match.endswith('*'):
                content = match.group(4)
                parts.append({
                    'text': content,
                    'bold': False,
                    'italic': True
                })
            elif full_match.startswith('_') and full_match.endswith('_'):
                content = match.group(6)
                parts.append({
                    'text': content,
                    'bold': False,
                    'italic': True
                })

            last_pos = match.end()

        # 添加剩余的普通文本
        if last_pos < len(text):
            parts.append({
                'text': text[last_pos:],
                'bold': False,
                'italic': False
            })

        # 如果没有匹配到任何格式，返回原始文本
        if not parts:
            parts.append({
                'text': text,
                'bold': False,
                'italic': False
            })

        return parts

    def convert(self, elements: List[dict]) -> Document:
        """转换元素列表为 DOCX 文档"""
        for element in elements:
            element_type = element.get('type')

            if element_type == 'heading':
                self.add_heading(element['content'], element['level'])
            elif element_type == 'paragraph':
                formatted_parts = element.get('formatted_parts')
                self.add_paragraph(element['content'], formatted_parts)
            elif element_type == 'code':
                self.add_code(element['content'], element.get('language', ''))
            elif element_type == 'image':
                self.add_image(element['url'], element.get('alt', ''))
            elif element_type == 'table':
                self.add_table(element['headers'], element['rows'])
            elif element_type == 'list':
                self.add_list(element['list_type'], element['items'])

        return self.document

    def save(self, output_path: str):
        """保存文档"""
        self.document.save(output_path)


def convert_md_to_docx(md_content: str, image_width: float = 5.0):
    try:
        # 解析 Markdown
        parser = MarkdownParser(md_content)
        elements = parser.parse()

        # 转换为 DOCX
        converter = DocxConverter(image_width=image_width)
        document = converter.convert(elements)

        # 将Document对象保存到字节流
        output_stream = BytesIO()
        document.save(output_stream)
        output_stream.seek(0)
        return 1, output_stream.getvalue()
    except Exception as e:
        return 0, None


class Md2docxTool(Tool):
    def _invoke(self, tool_parameters: dict[str, Any]) -> Generator[ToolInvokeMessage]:

        md_content = tool_parameters.get('content')
        output_name = tool_parameters.get('name') + ".docx"

        if not md_content:
            yield self.create_text_message("No markdown content provided.")
            return

        status, file_bytes = convert_md_to_docx(md_content)
        if status == 1:
            yield self.create_text_message(f"Document '{output_name}' generated successfully")
            yield self.create_blob_message(
                blob=file_bytes,
                meta={
                    "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "filename": output_name
                }
            )
        else:
            yield self.create_text_message("Error converting markdown to DOCX")
